from __future__ import annotations

import io
import re
import tempfile
from pathlib import Path
from typing import Any

from docx import Document
from googleapiclient.http import MediaIoBaseDownload
from openpyxl import load_workbook
from pypdf import PdfReader
from rich.console import Console

from .database import connect, list_extractable_files, upsert_file_text
from .drive_client import build_drive_service
from .search import search_files


console = Console()

GOOGLE_DOC = "application/vnd.google-apps.document"
GOOGLE_SLIDES = "application/vnd.google-apps.presentation"
GOOGLE_SHEETS = "application/vnd.google-apps.spreadsheet"

PDF = "application/pdf"
DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PPTX = "application/vnd.openxmlformats-officedocument.presentationml.presentation"
XLSX = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

TEXT = "text/plain"
CSV = "text/csv"
HTML = "text/html"

# Keep this conservative. You can add HTML back later if needed.
DEFAULT_EXTRACTABLE_MIMES = {
    GOOGLE_DOC,
    GOOGLE_SLIDES,
    GOOGLE_SHEETS,
    PDF,
    DOCX,
    XLSX,
    TEXT,
    CSV,
}


def clean_text(text: str, max_chars: int = 200_000) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()[:max_chars]


def _download_request_to_bytes(request) -> bytes:
    fh = io.BytesIO()
    downloader = MediaIoBaseDownload(fh, request)

    done = False
    while not done:
        _, done = downloader.next_chunk(num_retries=3)

    return fh.getvalue()


def _export_google_file(service, file_id: str, export_mime: str) -> bytes:
    request = service.files().export_media(
        fileId=file_id,
        mimeType=export_mime,
    )
    return _download_request_to_bytes(request)


def _download_binary_file(service, file_id: str) -> bytes:
    request = service.files().get_media(fileId=file_id)
    return _download_request_to_bytes(request)


def _extract_pdf_from_bytes(data: bytes, max_pages: int = 40) -> str:
    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []

    for page in reader.pages[:max_pages]:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue

    if len(reader.pages) > max_pages:
        parts.append(f"[PDF truncated after {max_pages} pages]")

    return clean_text("\n".join(parts))


def _extract_docx_from_bytes(data: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(data)
        tmp_path = Path(tmp.name)

    try:
        doc = Document(str(tmp_path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return clean_text("\n".join(parts))
    finally:
        tmp_path.unlink(missing_ok=True)


def _extract_xlsx_from_bytes(data: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
        tmp.write(data)
        tmp_path = Path(tmp.name)

    try:
        wb = load_workbook(str(tmp_path), read_only=True, data_only=True)
        parts: list[str] = []

        for ws in wb.worksheets[:10]:
            parts.append(f"Sheet: {ws.title}")
            row_count = 0

            for row in ws.iter_rows(values_only=True):
                values = [str(v) for v in row if v is not None and str(v).strip()]
                if values:
                    parts.append(" | ".join(values))
                    row_count += 1

                if row_count >= 200:
                    parts.append("[Sheet truncated after 200 non-empty rows]")
                    break

        return clean_text("\n".join(parts))
    finally:
        tmp_path.unlink(missing_ok=True)


def _decode_text_bytes(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return clean_text(data.decode(encoding, errors="ignore"))
        except Exception:
            continue
    return ""


def extract_text_for_file(account: str, file: dict[str, Any]) -> tuple[str, str]:
    service = build_drive_service(account)

    file_id = file["id"]
    mime = file.get("mime_type") or file.get("mimeType") or ""

    if mime == GOOGLE_DOC:
        data = _export_google_file(service, file_id, "text/plain")
        return clean_text(_decode_text_bytes(data)), "google_doc_text_export"

    if mime == GOOGLE_SLIDES:
        data = _export_google_file(service, file_id, "text/plain")
        return clean_text(_decode_text_bytes(data)), "google_slides_text_export"

    if mime == GOOGLE_SHEETS:
        data = _export_google_file(
            service,
            file_id,
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        return _extract_xlsx_from_bytes(data), "google_sheets_xlsx_export"

    if mime == PDF:
        data = _download_binary_file(service, file_id)
        return _extract_pdf_from_bytes(data), "pdf_download"

    if mime == DOCX:
        data = _download_binary_file(service, file_id)
        return _extract_docx_from_bytes(data), "docx_download"

    if mime == XLSX:
        data = _download_binary_file(service, file_id)
        return _extract_xlsx_from_bytes(data), "xlsx_download"

    if mime in {TEXT, CSV}:
        data = _download_binary_file(service, file_id)
        return _decode_text_bytes(data), "text_download"

    return "", "unsupported"


def already_extracted(file_id: str) -> bool:
    with connect() as conn:
        row = conn.execute(
            """
            SELECT extraction_status
            FROM file_text
            WHERE file_id = ?
            """,
            (file_id,),
        ).fetchone()

    return row is not None


def extract_files(
    account: str,
    files: list[dict[str, Any]],
    force: bool = False,
) -> tuple[int, int, int]:
    success = 0
    failed = 0
    skipped = 0

    total = len(files)

    for idx, file in enumerate(files, start=1):
        file_id = file["id"]
        name = file.get("name", "")
        mime = file.get("mime_type") or file.get("mimeType") or ""

        if mime not in DEFAULT_EXTRACTABLE_MIMES:
            skipped += 1
            console.print(
                f"[dim][{idx}/{total}] Skipping unsupported: {name} ({mime})[/dim]"
            )
            continue

        if not force and already_extracted(file_id):
            skipped += 1
            console.print(f"[dim][{idx}/{total}] Already extracted: {name}[/dim]")
            continue

        console.print(f"[cyan][{idx}/{total}] Extracting:[/cyan] {name}")
        console.print(f"   Type: {mime}")

        try:
            text, extractor = extract_text_for_file(account, file)

            if text.strip():
                upsert_file_text(
                    file_id=file_id,
                    extracted_text=text,
                    extraction_status="ok",
                    extractor=extractor,
                )
                success += 1
                console.print(
                    f"   [green]OK[/green] {len(text):,} chars via {extractor}"
                )
            else:
                upsert_file_text(
                    file_id=file_id,
                    extracted_text="",
                    extraction_status="empty",
                    extractor=extractor,
                )
                failed += 1
                console.print(f"   [yellow]EMPTY[/yellow] via {extractor}")

        except KeyboardInterrupt:
            console.print(
                "[yellow]Interrupted by user. Partial extraction progress was saved.[/yellow]"
            )
            raise

        except Exception as e:
            upsert_file_text(
                file_id=file_id,
                extracted_text="",
                extraction_status="error",
                extractor="unknown",
                error_message=f"{type(e).__name__}: {e}",
            )
            failed += 1
            console.print(f"   [red]ERROR[/red] {type(e).__name__}: {e}")

    return success, failed, skipped


def extract_pending(
    account: str,
    limit: int = 100,
    force: bool = False,
) -> tuple[int, int, int]:
    files = list_extractable_files(limit=limit, force=force)
    return extract_files(account=account, files=files, force=force)


def extract_for_query(
    account: str,
    query: str,
    limit: int = 50,
    force: bool = False,
) -> tuple[int, int, int]:
    hits = search_files(query, limit=limit)

    files: list[dict[str, Any]] = []
    seen: set[str] = set()

    for hit in hits:
        file_id = hit["id"]
        if file_id in seen:
            continue
        seen.add(file_id)
        files.append(hit)

    console.print(
        f"[bold]Extracting text for top {len(files)} local results for:[/bold] {query}"
    )
    return extract_files(account=account, files=files, force=force)
